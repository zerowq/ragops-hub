from __future__ import annotations

from pathlib import Path


class DocumentParser:
    def parse(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            from pypdf import PdfReader

            return "\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        if suffix == ".docx":
            from docx import Document

            document = Document(path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n\n".join(paragraphs)
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .txt, .md, .pdf, .docx")

